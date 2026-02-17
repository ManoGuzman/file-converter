"""
PDF to Word Converter Script.
"""

import sys
from pathlib import Path
from typing import Tuple

# Try to import pdf2docx
try:
    from pdf2docx import Converter

    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False
    print("Warning: pdf2docx not available")

# Try to import alternative conversion libraries
try:
    import pdfplumber
    from docx import Document

    ALTERNATIVE_AVAILABLE = True
except ImportError:
    ALTERNATIVE_AVAILABLE = False


# Constants
DEFAULT_PDF_DIR = "pdf"
DEFAULT_OUTPUT_DIR = "docs"
PDF_EXTENSION = ".pdf"
DOCX_EXTENSION = ".docx"
SEPARATOR_LINE = "=" * 60


def is_file_readable(path: Path) -> bool:
    """Check if a file is readable.

    Args:
        path: Path object to check

    Returns:
        True if file is readable, False otherwise
    """
    try:
        with open(path, "rb"):
            pass
        return True
    except (PermissionError, OSError):
        return False


def validate_path_exists(path: Path, path_type: str) -> None:
    """Validate that a path exists.

    Args:
        path: Path to validate
        path_type: Description of path type (for error messages)

    Raises:
        FileNotFoundError: If path doesn't exist
    """
    if not path.exists():
        raise FileNotFoundError(f"{path_type} not found: {path}")


def validate_is_file(path: Path) -> None:
    """Validate that a path is a file.

    Args:
        path: Path to validate

    Raises:
        ValueError: If path is not a file
    """
    if not path.is_file():
        raise ValueError(f"Path is not a file: {path}")


def validate_is_directory(path: Path) -> None:
    """Validate that a path is a directory.

    Args:
        path: Path to validate

    Raises:
        ValueError: If path is not a directory
    """
    if not path.is_dir():
        raise ValueError(f"Path is not a directory: {path}")


def validate_file_is_readable(path: Path) -> None:
    """Validate that a file is readable.

    Args:
        path: Path to validate

    Raises:
        PermissionError: If file is not readable
    """
    if not is_file_readable(path):
        raise PermissionError(f"File is not readable: {path}")


def validate_file_not_empty(path: Path) -> None:
    """Validate that a file is not empty.

    Args:
        path: Path to validate

    Raises:
        ValueError: If file is empty
    """
    if path.stat().st_size == 0:
        raise ValueError(f"File is empty: {path}")


def validate_file_has_extension(path: Path, expected_ext: str) -> None:
    """Validate file extension.

    Args:
        path: Path to validate
        expected_ext: Expected file extension (e.g., '.pdf')

    Raises:
        ValueError: If extension doesn't match
    """
    actual_ext = path.suffix.lower()
    if actual_ext != expected_ext:
        raise ValueError(
            f"File must have {expected_ext} extension (got '{actual_ext}'): {path}"
        )


def validate_pdf_input_file(pdf_file: str) -> Path:
    """Validate that the input PDF file exists and is readable.

    Args:
        pdf_file: Path to PDF file as string

    Returns:
        Validated Path object

    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If file is invalid
        PermissionError: If file is not readable
    """
    path = Path(pdf_file)

    validate_path_exists(path, "Input file")
    validate_is_file(path)
    validate_file_is_readable(path)
    validate_file_not_empty(path)
    validate_file_has_extension(path, PDF_EXTENSION)

    return path


def validate_input_directory(pdf_dir: str) -> Tuple[Path, list]:
    """Validate that the input directory exists and contains PDF files.

    Args:
        pdf_dir: Path to directory as string

    Returns:
        Tuple of (Path object, list of PDF files)

    Raises:
        FileNotFoundError: If directory doesn't exist
        ValueError: If directory is invalid or contains no PDFs
    """
    path = Path(pdf_dir)

    validate_path_exists(path, "Input directory")
    validate_is_directory(path)

    pdf_files = list(path.glob(f"*{PDF_EXTENSION}"))
    if not pdf_files:
        raise ValueError(f"No PDF files found in directory: {pdf_dir}")

    return path, pdf_files


def validate_output_directory_exists(output_dir: str) -> Path:
    """Validate or create the output directory.

    Args:
        output_dir: Path to output directory as string

    Returns:
        Validated Path object

    Raises:
        ValueError: If path exists but is not a directory
    """
    path = Path(output_dir)

    if not path.exists():
        print(f"Creating output directory: {output_dir}")
        path.mkdir(parents=True, exist_ok=True)
    else:
        validate_is_directory(path)

    return path


def validate_output_file(word_file: Path) -> Path:
    """Validate that the output path is writable.

    Args:
        word_file: Path to output file

    Returns:
        Validated Path object

    Raises:
        FileNotFoundError: If output directory doesn't exist
    """
    if word_file.exists():
        print(f"Warning: Output file will be overwritten: {word_file}")

    if not word_file.parent.exists():
        raise FileNotFoundError(f"Output directory does not exist: {word_file.parent}")

    if word_file.suffix.lower() != DOCX_EXTENSION:
        print(
            f"Warning: Output extension is '{word_file.suffix}', "
            f"expected '{DOCX_EXTENSION}'"
        )

    return word_file


def add_page_text_to_document(page, doc: Document) -> None:
    """Extract text from a PDF page and add to document.

    Args:
        page: pdfplumber page object
        doc: python-docx Document object
    """
    text = page.extract_text()
    if not text:
        return

    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)


def add_page_tables_to_document(page, doc: Document) -> None:
    """Extract tables from a PDF page and add to document.

    Args:
        page: pdfplumber page object
        doc: python-docx Document object
    """
    tables = page.extract_tables()

    for table in tables:
        if not table:
            continue

        max_cols = max((len(row) for row in table if row), default=0)
        if max_cols == 0:
            continue

        doc_table = doc.add_table(rows=len(table), cols=max_cols)

        for i, row in enumerate(table):
            for j in range(max_cols):
                cell_value = row[j] if j < len(row) else ""
                if cell_value:
                    doc_table.rows[i].cells[j].text = str(cell_value)


def convert_using_pdfplumber(pdf_file: Path, word_file: Path) -> None:
    """Convert PDF to Word using pdfplumber + python-docx.

    Args:
        pdf_file: Path to input PDF
        word_file: Path to output DOCX

    Raises:
        RuntimeError: If conversion fails or libraries not available
    """
    if not ALTERNATIVE_AVAILABLE:
        raise RuntimeError(
            "Alternative conversion libraries not available. "
            "Install pdfplumber and python-docx."
        )

    print(f"Converting {pdf_file.name} to {word_file.name} using alternative method...")

    try:
        doc = Document()

        with pdfplumber.open(str(pdf_file)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                if page_num > 1:
                    doc.add_page_break()

                add_page_text_to_document(page, doc)
                add_page_tables_to_document(page, doc)

        doc.save(str(word_file))
        print(f"✓ Successfully converted to {word_file}")

    except Exception as e:
        raise RuntimeError(
            f"Alternative conversion failed for {pdf_file.name}: {str(e)}"
        ) from e


def try_convert_using_pdf2docx(pdf_file: Path, word_file: Path) -> bool:
    """Try to convert using pdf2docx.

    Args:
        pdf_file: Path to input PDF
        word_file: Path to output DOCX

    Returns:
        True if conversion succeeded, False if compatibility issue detected

    Raises:
        RuntimeError: If conversion fails (non-compatibility errors)
    """
    cv = None
    try:
        cv = Converter(str(pdf_file))
        print(f"Converting {pdf_file.name} to {word_file.name}...")
        cv.convert(str(word_file), start=0)
        print(f"✓ Successfully converted to {word_file}")
        return True

    except AttributeError as e:
        # Compatibility error with PyMuPDF - fall back to alternative
        print(f"⚠ pdf2docx compatibility issue: {str(e)}")
        print("⚠ Falling back to alternative conversion method...")
        return False

    except (OSError, RuntimeError, ValueError) as e:
        raise RuntimeError(f"Conversion failed for {pdf_file.name}: {str(e)}") from e

    finally:
        if cv:
            try:
                cv.close()
            except Exception:  # pylint: disable=broad-except
                pass  # Ignore cleanup errors


def convert_pdf_to_word(pdf_file: Path, word_file: Path) -> None:
    """Convert a PDF file to a Word document.

    Tries pdf2docx first, falls back to pdfplumber + python-docx if needed.

    Args:
        pdf_file: Path to input PDF
        word_file: Path to output DOCX

    Raises:
        RuntimeError: If conversion fails
    """
    if PDF2DOCX_AVAILABLE and try_convert_using_pdf2docx(pdf_file, word_file):
        return

    convert_using_pdfplumber(pdf_file, word_file)


def process_single_pdf_file(input_pdf: str, output_dir: str) -> None:
    """Convert a single PDF file to the output directory.

    Args:
        input_pdf: Path to input PDF file
        output_dir: Path to output directory

    Raises:
        Various exceptions from validation and conversion
    """
    pdf_path = validate_pdf_input_file(input_pdf)
    output_path = validate_output_directory_exists(output_dir)

    word_filename = pdf_path.stem + DOCX_EXTENSION
    word_path = output_path / word_filename

    validate_output_file(word_path)
    convert_pdf_to_word(pdf_path, word_path)


def process_pdf_directory(input_dir: str, output_dir: str) -> None:
    """Convert all PDF files in a directory to the output directory.

    Args:
        input_dir: Path to input directory
        output_dir: Path to output directory
    """
    _, pdf_files = validate_input_directory(input_dir)
    output_path = validate_output_directory_exists(output_dir)

    total = len(pdf_files)
    print(f"\nFound {total} PDF file(s) to convert\n")

    converted = 0
    failed = 0

    for i, pdf_file in enumerate(pdf_files, 1):
        try:
            word_filename = pdf_file.stem + DOCX_EXTENSION
            word_path = output_path / word_filename

            print(f"[{i}/{total}] ", end="")
            convert_pdf_to_word(pdf_file, word_path)
            converted += 1

        except (RuntimeError, ValueError, OSError) as e:
            print(f"✗ Failed: {e}", file=sys.stderr)
            failed += 1

    print_conversion_summary(converted, failed)


def print_conversion_summary(converted: int, failed: int) -> None:
    """Print summary of batch conversion results.

    Args:
        converted: Number of successful conversions
        failed: Number of failed conversions
    """
    print(f"\n{SEPARATOR_LINE}")
    print(f"Conversion complete: {converted} successful, {failed} failed")
    print(SEPARATOR_LINE)


def print_usage() -> None:
    """Print usage information."""
    usage_text = """PDF to Word Converter

Usage:
  Mode 1 - Convert single file:
    python src/convert_pdf_to_word.py -f <input_pdf>
    Output: convert_files/docs/

  Mode 2 - Convert all PDFs in directory:
    python src/convert_pdf_to_word.py -d
    Input: convert_files/pdf/
    Output: convert_files/docs/

Examples:
  python src/convert_pdf_to_word.py -f "C:\\path\\to\\file.pdf"
  python src/convert_pdf_to_word.py -d"""

    print(usage_text)


def check_dependencies() -> None:
    """Check if required conversion libraries are available.

    Raises:
        SystemExit: If no conversion libraries are available
    """
    if not PDF2DOCX_AVAILABLE and not ALTERNATIVE_AVAILABLE:
        print("Error: No PDF conversion libraries available!")
        print("\nPlease install required dependencies:")
        print("  pip install pymupdf pdfplumber python-docx")
        sys.exit(1)


def get_default_directories() -> Tuple[Path, Path]:
    """Get default input and output directories.

    Returns:
        Tuple of (input_dir, output_dir) as Path objects
    """
    script_dir = Path(__file__).parent.parent
    return script_dir / DEFAULT_PDF_DIR, script_dir / DEFAULT_OUTPUT_DIR


def handle_file_mode(default_output_dir: Path) -> None:
    """Handle single file conversion mode.

    Args:
        default_output_dir: Default output directory
    """
    if len(sys.argv) != 3:
        print("Error: Mode -f requires an input PDF file")
        print("\nUsage: python src/convert_pdf_to_word.py -f <input_pdf>")
        sys.exit(1)

    input_pdf = sys.argv[2]
    process_single_pdf_file(input_pdf, str(default_output_dir))


def handle_directory_mode(default_input_dir: Path, default_output_dir: Path) -> None:
    """Handle directory conversion mode.

    Args:
        default_input_dir: Default input directory
        default_output_dir: Default output directory
    """
    process_pdf_directory(str(default_input_dir), str(default_output_dir))


def main() -> None:
    """Main entry point for the script."""
    check_dependencies()

    if len(sys.argv) < 2:
        print_usage()
        sys.exit(1)

    default_input_dir, default_output_dir = get_default_directories()
    mode = sys.argv[1]

    try:
        match mode:
            case "-f" | "--file":
                handle_file_mode(default_output_dir)
            case "-d" | "--directory":
                handle_directory_mode(default_input_dir, default_output_dir)
            case _:
                print(f"Error: Unknown mode '{mode}'")
                print_usage()
                sys.exit(1)

    except (FileNotFoundError, ValueError, RuntimeError, PermissionError) as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
    except KeyboardInterrupt:
        print("\nConversion cancelled by user", file=sys.stderr)
        sys.exit(130)


if __name__ == "__main__":
    main()
